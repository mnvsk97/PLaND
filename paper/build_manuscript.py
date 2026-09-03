#!/usr/bin/env python3
"""Build the PLaND SAS Publishers submission manuscript from Markdown.

The generated DOCX follows the published journal article format: A4, compact
Times New Roman typography, full-width front matter, two-column body text,
spanning tables/figures, running headers, and bordered footers. The Markdown
remains the source of truth.
"""

from __future__ import annotations

import argparse
import copy
import html
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image


FONT = "Times New Roman"
CONTENT_WIDTH_DXA = 9890  # A4 width minus two 0.7-inch margins.
FIGURE_MARKERS = {
    "<!-- architecture-diagram -->": "architecture",
    "<!-- evolution-path-diagram -->": "evolution-path",
    "<!-- evolution-diagram -->": "evolution-loop",
}
FIGURE_WIDTHS = {
    "architecture": 3.20,
    "evolution-path": 5.00,
    "evolution-loop": 2.75,
}
REQUIRED_SECTIONS = (
    "Abstract",
    "Introduction",
    "Material and Methods",
    "Results",
    "Discussion",
    "Conclusion",
    "References",
)


def set_run_font(run, *, size: float = 9, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 4,
                          keep_with_next: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.0
    fmt.keep_with_next = keep_with_next


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    properties.append(size)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*([^*]+)\*\*|`([^`]+)`|(?<!\*)\*([^*]+)\*(?!\*)|(https?://\S+))"
)


def add_inline(paragraph, value: str) -> None:
    """Add a small, deterministic subset of Markdown inline formatting."""
    cursor = 0
    for match in INLINE_PATTERN.finditer(value):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(html.unescape(value[cursor:match.start()])))
        token = match.group(0)
        if match.group(2) and match.group(3):
            add_hyperlink(paragraph, match.group(2), match.group(3))
        elif match.group(4):
            set_run_font(paragraph.add_run(match.group(4)), bold=True)
        elif match.group(5):
            run = paragraph.add_run(match.group(5))
            set_run_font(run)
            run.font.name = "Courier New"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
        elif match.group(6):
            set_run_font(paragraph.add_run(match.group(6)), italic=True)
        elif match.group(7):
            url = token.rstrip(".,;:")
            add_hyperlink(paragraph, url, url)
            if len(url) != len(token):
                set_run_font(paragraph.add_run(token[len(url):]))
        cursor = match.end()
    if cursor < len(value):
        set_run_font(paragraph.add_run(html.unescape(value[cursor:])))


def set_cell_margins(cell, *, top: int = 80, start: int = 100,
                     bottom: int = 80, end: int = 100) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "557030")


def set_paragraph_rule(paragraph, color: str = "557030") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_columns(section, count: int) -> None:
    properties = section._sectPr
    columns = properties.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        properties.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), "340")


def add_continuous_section(document: Document, columns: int):
    section = document.add_section(WD_SECTION.CONTINUOUS)
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.3)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    set_columns(section, columns)
    return section


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    scores = []
    for index in range(columns):
        longest = max(len(row[index]) for row in rows)
        scores.append(max(10, min(longest, 48)))
    # Keep row labels readable even when another column contains long prose.
    scores[0] = max(scores[0], 18)
    total = sum(scores)
    widths = [round(CONTENT_WIDTH_DXA * score / total) for score in scores]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(document: Document, rows: list[list[str]]) -> None:
    widths = table_widths(rows)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "0")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for row_index, (word_row, docx_row) in enumerate(zip(rows, table.rows)):
        for column_index, (value, cell) in enumerate(zip(word_row, docx_row.cells)):
            set_cell_width(cell, widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0)
            if column_index > 0 and re.fullmatch(r"[\d.,%+\-\s]+", value):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=8.2, bold=(row_index == 0))
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=2)


def set_image_alt_text(inline_shape, description: str) -> None:
    drawing = inline_shape._inline
    properties = drawing.docPr
    properties.set("descr", description)


def prepare_jpeg(figures_dir: Path, name: str) -> Path:
    jpg = figures_dir / f"{name}.jpg"
    svg = figures_dir / f"{name}.svg"
    if not svg.exists():
        raise FileNotFoundError(f"Missing figure source: {svg}")
    svg_text = svg.read_text(encoding="utf-8")
    size_match = re.search(r'<svg\s+width="([0-9.]+)"\s+height="([0-9.]+)"', svg_text)
    expected_ratio = (float(size_match.group(1)) / float(size_match.group(2))) if size_match else None
    if jpg.exists() and jpg.stat().st_mtime >= svg.stat().st_mtime:
        with Image.open(jpg) as existing:
            actual_ratio = existing.width / existing.height
        if expected_ratio is None or abs(actual_ratio - expected_ratio) <= 0.01:
            return jpg

    with tempfile.TemporaryDirectory(prefix="pland-figure-") as directory:
        temp = Path(directory)
        png = temp / f"{name}.png"
        try:
            import cairosvg  # type: ignore

            cairosvg.svg2png(url=str(svg), write_to=str(png), output_width=2400)
        except ImportError:
            qlmanage = shutil.which("qlmanage")
            if not qlmanage:
                raise RuntimeError(
                    "SVG conversion needs the free cairosvg package or macOS qlmanage"
                )
            subprocess.run(
                [qlmanage, "-t", "-s", "2400", "-o", str(temp), str(svg)],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            generated = temp / f"{svg.name}.png"
            if not generated.exists():
                raise RuntimeError(f"qlmanage did not create {generated}")
            generated.replace(png)
        with Image.open(png) as image:
            rgb = Image.new("RGB", image.size, "white")
            if image.mode == "RGBA":
                rgb.paste(image.convert("RGBA"), mask=image.getchannel("A"))
            else:
                rgb.paste(image.convert("RGB"))
            if name == "evolution-loop":
                padded = Image.new("RGB", (rgb.width, rgb.height + 48), "white")
                padded.paste(rgb, (0, 0))
                rgb = padded
            jpg.parent.mkdir(parents=True, exist_ok=True)
            rgb.save(jpg, "JPEG", quality=95, dpi=(300, 300), optimize=True)
    return jpg


def add_figure(document: Document, figures_dir: Path, name: str) -> None:
    image = prepare_jpeg(figures_dir, name)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=3, after=6, keep_with_next=True)
    shape = paragraph.add_run().add_picture(
        str(image), width=Inches(FIGURE_WIDTHS.get(name, 5.15))
    )
    set_image_alt_text(shape, name.replace("-", " ").title())


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, before, after in (("Heading 1", 10, 3), ("Heading 2", 8, 2)):
        style = document.styles[style_name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5 if style_name == "Heading 1" else 9.2)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 112, 192) if style_name == "Heading 1" else RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(9)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.0


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True
    set_columns(section, 1)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0)
    set_run_font(header.add_run("Maddipatla Naga Venkata Sai Krishna & Asit Kumar Sahoo, Sch J Eng Tech"), size=7)
    set_paragraph_rule(header)

    first_header = section.first_page_header.paragraphs[0]
    set_paragraph_spacing(first_header, after=0)

    footer = section.footer
    base = footer.paragraphs[0]
    base._element.getparent().remove(base._element)
    table = footer.add_table(rows=1, cols=2, width=Inches(6.86))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_width(table.cell(0, 0), 8070)
    set_cell_width(table.cell(0, 1), 1820)
    left = table.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(left, after=0)
    set_run_font(left.add_run("PLaND | Submission manuscript"), size=7)
    right_cell = table.cell(0, 1)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "DCE8C6")
    right_cell._tc.get_or_add_tcPr().append(shading)
    page = right_cell.paragraphs[0]
    page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(page, after=0)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = page.add_run()
    set_run_font(run, size=7)
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    first_footer = section.first_page_footer
    first_base = first_footer.paragraphs[0]
    first_base._element.getparent().remove(first_base._element)
    first_footer._element.append(copy.deepcopy(table._tbl))


def validate_source(markdown: str, *, allow_pending: bool) -> None:
    headings = re.findall(r"^## (.+)$", markdown, flags=re.MULTILINE)
    missing = [name for name in REQUIRED_SECTIONS if name not in headings]
    if missing:
        raise ValueError(f"Missing required manuscript sections: {', '.join(missing)}")
    required_positions = [headings.index(name) for name in REQUIRED_SECTIONS]
    if required_positions != sorted(required_positions):
        raise ValueError("Required manuscript sections are not in SAS research-article order")
    if not allow_pending and "RESULT_PENDING:" in markdown:
        raise ValueError(
            "Unresolved RESULT_PENDING markers remain. Insert saved results or pass "
            "--allow-pending-results for an internal draft."
        )
    abstract_match = re.search(r"^## Abstract\s+(.+?)(?=^\*\*Keywords:)", markdown, re.S | re.M)
    if not abstract_match:
        raise ValueError("Abstract or keywords block not found")
    abstract_words = re.findall(r"\b[\w'-]+\b", abstract_match.group(1))
    if len(abstract_words) > 300:
        raise ValueError(f"Abstract has {len(abstract_words)} words; maximum is 300")
    keywords_match = re.search(r"^\*\*Keywords:\*\*\s*(.+)$", markdown, re.M)
    keywords = [item.strip() for item in keywords_match.group(1).split(",")]
    if not 5 <= len(keywords) <= 10:
        raise ValueError(f"Expected 5-10 keywords, found {len(keywords)}")


def flush_body(document: Document, lines: list[str], *, reference: bool = False) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines)
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=1 if reference else 4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.2)
    if reference:
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
    add_inline(paragraph, text)
    lines.clear()


def build(markdown: str, destination: Path, figures_dir: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_document(document)
    document.core_properties.title = "PLaND - Path to Least Non Determinism"
    document.core_properties.subject = "Controlled evolution of agent skills"
    document.core_properties.keywords = "agent skills, deterministic workflows, evaluation"
    document.core_properties.author = "Maddipatla Naga Venkata Sai Krishna; Asit Kumar Sahoo"

    masthead = document.add_paragraph()
    set_paragraph_spacing(masthead, after=0, keep_with_next=True)
    set_run_font(masthead.add_run("Scholars Journal of Engineering and Technology"), size=18)
    masthead.runs[0].font.color.rgb = RGBColor(0, 74, 147)
    detail = document.add_paragraph()
    set_paragraph_spacing(detail, after=8, keep_with_next=True)
    set_run_font(detail.add_run("Abbreviated Key Title: Sch J Eng Tech\nJournal homepage: https://saspublishers.com/journal/sjet/home"), size=7.5)
    set_paragraph_rule(detail)

    lines = markdown.splitlines()
    body: list[str] = []
    in_code = False
    code_lines: list[str] = []
    in_references = False
    first_heading = True
    columns_started = False
    span_kind = None
    index = 0

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            flush_body(document, body, reference=in_references)
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(paragraph, before=2, after=5)
                for line_index, line in enumerate(code_lines):
                    run = paragraph.add_run(line)
                    set_run_font(run, size=10)
                    if line_index < len(code_lines) - 1:
                        run.add_break()
                code_lines.clear()
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped in FIGURE_MARKERS:
            flush_body(document, body, reference=in_references)
            if columns_started:
                add_continuous_section(document, 1)
                span_kind = "figure"
            add_figure(document, figures_dir, FIGURE_MARKERS[stripped])
            index += 1
            continue

        if stripped.startswith("<!--") and stripped.endswith("-->"):
            index += 1
            continue

        if raw.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            flush_body(document, body, reference=in_references)
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            parsed_rows = []
            for line in table_lines:
                cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    continue
                parsed_rows.append(cells)
            add_table(document, parsed_rows)
            if columns_started and span_kind == "table":
                add_continuous_section(document, 2)
                span_kind = None
            continue

        list_match = re.match(r"^(\d+)\.\s+(.+)$", raw)
        if raw.startswith("- ") or list_match:
            flush_body(document, body, reference=in_references)
            if list_match:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.30)
                paragraph.paragraph_format.first_line_indent = Inches(-0.20)
                content = f"{list_match.group(1)}. {list_match.group(2)}"
            else:
                paragraph = document.add_paragraph(style="List Bullet")
                content = raw[2:]
            set_paragraph_spacing(paragraph, after=2)
            add_inline(paragraph, content)
            index += 1
            continue

        if raw.startswith("> "):
            flush_body(document, body, reference=in_references)
            quote_lines = []
            while index < len(lines) and lines[index].startswith("> "):
                quote_lines.append(lines[index][2:])
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.paragraph_format.right_indent = Inches(0.35)
            set_paragraph_spacing(paragraph, before=3, after=5)
            run = paragraph.add_run(" ".join(quote_lines))
            set_run_font(run, italic=True)
            continue

        if raw.startswith("# "):
            flush_body(document, body, reference=in_references)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, before=0, after=3, keep_with_next=True)
            run = paragraph.add_run(raw[2:])
            set_run_font(run, size=16, bold=False)
            run.font.color.rgb = RGBColor(0, 74, 147)
            set_paragraph_rule(paragraph, "D21F26")
            index += 1
            continue

        if raw.startswith("## "):
            flush_body(document, body, reference=in_references)
            heading_text = raw[3:]
            paragraph = document.add_paragraph(style="Heading 1")
            if heading_text == "Abstract":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(heading_text.upper())
            set_run_font(run, bold=True)
            if heading_text == "Abstract":
                shade_paragraph(paragraph, "002060")
                run.font.color.rgb = RGBColor(255, 255, 255)
            in_references = heading_text == "References"
            first_heading = False
            index += 1
            continue

        if raw.startswith("### "):
            flush_body(document, body, reference=in_references)
            paragraph = document.add_paragraph(style="Heading 2")
            run = paragraph.add_run(raw[4:])
            set_run_font(run, bold=True)
            index += 1
            continue

        if first_heading and stripped.startswith("**Maddipatla Naga Venkata Sai Krishna**"):
            flush_body(document, body)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, after=3, keep_with_next=True)
            add_inline(paragraph, stripped)
            index += 1
            continue

        if first_heading and (stripped.startswith("*Author affiliations") or stripped.startswith("*Affiliations")):
            flush_body(document, body)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, after=8, keep_with_next=True)
            run = paragraph.add_run(stripped.strip("*"))
            set_run_font(run, italic=True)
            index += 1
            continue

        if stripped.startswith("**Figure "):
            flush_body(document, body, reference=in_references)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, after=6, keep_with_next=False)
            add_inline(paragraph, stripped)
            if columns_started and span_kind == "figure":
                add_continuous_section(document, 2)
                span_kind = None
            index += 1
            continue

        if stripped.startswith("**Table "):
            flush_body(document, body, reference=in_references)
            if columns_started:
                add_continuous_section(document, 1)
                span_kind = "table"
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, before=4, after=2, keep_with_next=True)
            add_inline(paragraph, stripped)
            index += 1
            continue

        if stripped.startswith("**Keywords:**"):
            flush_body(document, body, reference=in_references)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(0)
            set_paragraph_spacing(paragraph, after=4)
            shade_paragraph(paragraph, "F1F5E8")
            set_paragraph_rule(paragraph)
            add_inline(paragraph, stripped)
            add_continuous_section(document, 2)
            columns_started = True
            index += 1
            continue

        if not stripped:
            flush_body(document, body, reference=in_references)
        else:
            body.append(raw)
        index += 1

    flush_body(document, body, reference=in_references)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    default_source = Path(__file__).with_name("PAPER.md")
    parser.add_argument("source", nargs="?", type=Path, default=default_source)
    parser.add_argument("destination", nargs="?", type=Path,
                        default=Path("output/docx/PLAND_SAS_SUBMISSION.docx"))
    parser.add_argument("--figures-dir", type=Path,
                        default=Path(__file__).with_name("figures"))
    parser.add_argument("--allow-pending-results", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    markdown = args.source.read_text(encoding="utf-8")
    validate_source(markdown, allow_pending=args.allow_pending_results)
    build(markdown, args.destination, args.figures_dir)
    print(args.destination.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
