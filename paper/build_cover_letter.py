#!/usr/bin/env python3
"""Build the prepared-only SJET submission cover letter."""

from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


FONT = "Times New Roman"
NAVY = RGBColor(0, 74, 147)
MUTED = RGBColor(90, 90, 90)


def font(run, size: float, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def spacing(paragraph, before: float = 0, after: float = 7, line: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def rule(paragraph, color: str = "557030") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def paragraph(document: Document, text: str, *, size: float = 11, bold: bool = False,
              color: RGBColor | None = None, before: float = 0, after: float = 7,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    item = document.add_paragraph()
    item.alignment = align
    spacing(item, before, after)
    font(item.add_run(text), size, bold=bold, color=color)
    return item


def build(destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    spacing(header, after=0)
    font(header.add_run("PLaND | Original Research Article"), 8.5, color=MUTED)
    rule(header)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(footer, after=0)
    font(footer.add_run("Prepared for SJET | Not submitted"), 8, color=MUTED)

    paragraph(document, "Scholars Journal of Engineering and Technology", size=11,
              bold=True, color=NAVY, after=2)
    paragraph(document, "Editorial Office", size=10.5, color=MUTED, after=16)
    title = paragraph(document, "COVER LETTER", size=20, color=NAVY, after=3)
    rule(title, "D21F26")
    paragraph(document, "Submission type: Original Research Article", size=10,
              bold=True, color=MUTED, after=14)
    paragraph(document, "Dear Editorial Team,", after=10)
    paragraph(document,
              "Please consider our manuscript, “PLaND - Path to Least Non Determinism,” "
              "for publication as an Original Research Article in the Scholars Journal of "
              "Engineering and Technology.")
    paragraph(document,
              "The manuscript presents PLaND as an evaluation-driven methodology for moving "
              "stable workflow steps from natural language into tested deterministic "
              "implementations while retaining language-model judgment where evidence requires "
              "it. A three-seed paired study provides direct run-to-run evidence: LEDGAR is the "
              "repeated positive case, while CFPB and SpamAssassin preserve rejection and safety "
              "evidence under the same frozen gate structure.")
    paragraph(document,
              "The complete experiment evidence, manifests, exact commands, case-level outputs, "
              "and reproducibility metadata are preserved in the same Git commit as this "
              "manuscript, on branch codex/pland-variance-study. An arXiv-ready source bundle "
              "has been prepared locally but has "
              "not been uploaded or submitted; no arXiv identifier exists.")
    paragraph(document,
              "Both authors are affiliated as Independent Researcher, San Francisco, CA, USA. "
              "The work received no external funding. The authors declare no conflicts of "
              "interest, no acknowledgements, and no ORCID identifiers.")
    paragraph(document, "Thank you for your consideration.", before=3, after=12)
    paragraph(document, "Sincerely,", after=8)
    paragraph(document, "Maddipatla Naga Venkata Sai Krishna", bold=True, after=1)
    paragraph(document, "Corresponding author", size=10, color=MUTED, after=1)
    paragraph(document, "Independent Researcher, San Francisco, CA, USA", size=10, after=1)
    paragraph(document, "mnvsk97@gmail.com", size=10, color=NAVY, after=8)
    paragraph(document, "Asit Kumar Sahoo", bold=True, after=1)
    paragraph(document, "Independent Researcher, San Francisco, CA, USA", size=10, after=0)

    document.core_properties.title = "SJET cover letter - PLaND"
    document.core_properties.subject = "Prepared Original Research Article submission cover letter"
    document.core_properties.author = "Maddipatla Naga Venkata Sai Krishna; Asit Kumar Sahoo"
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def convert_pdf(docx: Path, pdf: Path) -> None:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice is required for cover-letter PDF export")
    with tempfile.TemporaryDirectory(prefix="pland-cover-letter-") as temporary:
        subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", temporary,
                        str(docx)], check=True)
        generated = Path(temporary) / f"{docx.stem}.pdf"
        shutil.copy2(generated, pdf)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=Path("output/paper/PLAND_SJET_COVER_LETTER.docx"))
    parser.add_argument("--pdf", type=Path, default=Path("output/paper/PLAND_SJET_COVER_LETTER.pdf"))
    args = parser.parse_args()
    build(args.docx)
    convert_pdf(args.docx.resolve(), args.pdf.resolve())
    print(args.docx.resolve())
    print(args.pdf.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
